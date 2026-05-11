#!/usr/bin/env python3
"""Generate one Word (.docx) handout per attendee.

Reads everything from `terraform output` (so it stays in sync with the
live landing zone) and writes one .docx per attendee containing a
two-column table. The Name / Email / TAP rows are intentionally left
blank — fill them in inside Word before sending.

Usage:
    /tmp/docxenv/bin/python scripts/generate-handouts-docx.py \\
        --out ~/handouts-docx \\
        --count 10

Requires:
    python-docx (installed in /tmp/docxenv from the facilitator setup).
"""

import argparse
import json
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt, RGBColor


# ----- helpers ---------------------------------------------------------------

def tf_output(name: str, raw: bool = True, json_mode: bool = False) -> str:
    """Read a single `terraform output` value from infra/."""
    args = ["terraform", "-chdir=infra", "output"]
    if raw:
        args.append("-raw")
    elif json_mode:
        args.append("-json")
    args.append(name)
    res = subprocess.run(args, capture_output=True, text=True, check=True)
    return res.stdout.strip()


def az_apim_key(rg: str, apim: str, sid: str, subscription_id: str = "") -> str:
    """Pull the per-attendee APIM subscription primary key.

    Tries `az apim subscription show` first (newer az CLI), then falls
    back to the Management API `listSecrets` endpoint via `az rest` —
    same pattern as scripts/print-attendee-handout.sh.
    """
    res = subprocess.run(
        [
            "az", "apim", "subscription", "show",
            "--resource-group", rg,
            "--service-name", apim,
            "--sid", sid,
            "--query", "primaryKey",
            "-o", "tsv",
        ],
        capture_output=True, text=True, check=False,
    )
    if res.returncode == 0 and res.stdout.strip():
        return res.stdout.strip()

    if subscription_id:
        res = subprocess.run(
            [
                "az", "rest", "--method", "post",
                "--url", f"{subscription_id}/listSecrets?api-version=2022-08-01",
                "--query", "primaryKey", "-o", "tsv",
            ],
            capture_output=True, text=True, check=False,
        )
        if res.returncode == 0 and res.stdout.strip():
            return res.stdout.strip()
    return ""


# ----- docx building ---------------------------------------------------------

def write_handout_docx(path: Path, ns: str, ctx: dict) -> None:
    """Render one attendee's handout as a Word document with a table."""
    doc = Document()

    # Title
    title = doc.add_heading(f"AI Gateway Workshop — {ns}", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x16, 0x4B, 0x7A)

    doc.add_paragraph(
        "Welcome! Below are your connection details. The subscription "
        "key is sensitive — treat it like a password."
    )

    # ---- Identity (blank, to be filled in Word) ----------------------------
    doc.add_heading("Your identity", level=2)
    identity = doc.add_table(rows=4, cols=2)
    identity.style = "Light Grid Accent 1"
    identity.autofit = False
    identity.columns[0].width = Cm(5)
    identity.columns[1].width = Cm(11)
    for i, (k, v) in enumerate([
        ("Name",            ""),
        ("Email",           ""),
        ("Temporary Access Pass (TAP)", ""),
        ("Attendee number", ns),
    ]):
        identity.rows[i].cells[0].text = k
        identity.rows[i].cells[1].text = v

    # ---- Workshop connection ----------------------------------------------
    doc.add_heading("Workshop connection", level=2)
    conn = doc.add_table(rows=0, cols=2)
    conn.style = "Light Grid Accent 1"
    conn.autofit = False
    fields = [
        ("Namespace",            ns),
        ("APIM Gateway URL",     ctx["APIM_GATEWAY"]),
        ("APIM Subscription key (Ocp-Apim-Subscription-Key)", ctx["APIM_KEY"]),
        ("AOAI endpoint",        ctx["AOAI"]),
        ("AOAI chat deployment", ctx["GPT_DEPLOY"]),
        ("AOAI embedding deployment", ctx["EMB_DEPLOY"]),
        ("Cosmos endpoint",      ctx["COSMOS"]),
        ("Cosmos container",     f"state-{ns} (partition key: /sessionId)"),
        ("AI Search endpoint",   ctx["SEARCH"]),
        ("Key Vault URI",        ctx["KV_URI"]),
        ("AKS resource group",   ctx["RG"]),
        ("AKS cluster name",     ctx["AKS"]),
        ("APIM developer portal", ctx["APIM_PORTAL"]),
        ("App Insights connection string", ctx["APPI_CONN"]),
    ]
    for k, v in fields:
        row = conn.add_row()
        row.cells[0].text = k
        row.cells[1].text = v
    for col, width in [(0, Cm(5)), (1, Cm(11))]:
        for row in conn.rows:
            row.cells[col].width = width

    # ---- How to use --------------------------------------------------------
    doc.add_heading("How to use this handout", level=2)

    doc.add_paragraph("1. Sign in to Azure (use the TAP above for your first sign-in if provided):")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("    az login\n").font.name = "Consolas"
    p.add_run("    az account set --subscription <workshop-subscription-id>").font.name = "Consolas"

    doc.add_paragraph("2. Connect to the workshop AKS cluster:")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        f'    az aks get-credentials -g {ctx["RG"]} -n {ctx["AKS"]} --overwrite-existing\n'
    ).font.name = "Consolas"
    p.add_run(
        f'    kubectl config set-context --current --namespace={ns}'
    ).font.name = "Consolas"

    doc.add_paragraph("3. Export the two env vars every lab needs:")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        f'    export APIM_GATEWAY_URL="{ctx["APIM_GATEWAY"]}"\n'
    ).font.name = "Consolas"
    p.add_run(
        f'    export APIM_KEY="{ctx["APIM_KEY"]}"'
    ).font.name = "Consolas"

    doc.add_paragraph("4. Smoke-test the gateway:")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        f'    curl -s "${{APIM_GATEWAY_URL}}/openai/deployments/'
        f'{ctx["GPT_DEPLOY"]}/chat/completions?api-version=2024-10-21" \\\n'
        f'      -H "Ocp-Apim-Subscription-Key: ${{APIM_KEY}}" \\\n'
        f'      -H "Content-Type: application/json" \\\n'
        f'      -d \'{{"messages":[{{"role":"user","content":"hello"}}]}}\''
    ).font.name = "Consolas"

    doc.add_paragraph(
        "5. Verify the gateway policies any time after M1:"
    )
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("    ./scripts/verify-policies.sh         # after M1\n").font.name = "Consolas"
    p.add_run("    ./scripts/verify-policies.sh --m2    # after M2").font.name = "Consolas"

    doc.add_paragraph(
        "Workshop materials: "
        "https://adindabudi.github.io/azure-hybrid-ai-platform-workshop/"
    )

    # Code-style font for monospace cells
    for table in (identity, conn):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.save(path)


# ----- main ------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", type=Path, required=True,
                    help="Output directory (e.g. ~/handouts-docx).")
    ap.add_argument("--count", type=int, default=10,
                    help="How many attendees (defaults to 10).")
    args = ap.parse_args()

    out_dir = args.out.expanduser().resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_dir.chmod(0o700)

    repo = Path(__file__).resolve().parents[1]
    os.chdir(repo)

    print(f"==> Reading Terraform outputs from {repo}/infra ...")
    ctx_static = {
        "RG":          tf_output("resource_group_name"),
        "AKS":         tf_output("aks_name"),
        "APIM":        tf_output("apim_name"),
        "APIM_GATEWAY": tf_output("apim_gateway_url"),
        "APIM_PORTAL": tf_output("apim_developer_portal_url"),
        "AOAI":        tf_output("aoai_endpoint"),
        "GPT_DEPLOY":  tf_output("aoai_gpt_4o_mini_deployment"),
        "EMB_DEPLOY":  tf_output("aoai_embedding_deployment"),
        "COSMOS":      tf_output("cosmos_endpoint"),
        "SEARCH":      tf_output("search_endpoint"),
        "KV_URI":      tf_output("key_vault_uri"),
        "APPI_CONN":   tf_output("application_insights_connection_string"),
    }

    # The attendee_handout output is marked sensitive — read via -json
    # and parse to extract each attendee's APIM subscription resource ID
    # for the listSecrets fallback.
    handout_json = subprocess.run(
        ["terraform", "-chdir=infra", "output", "-json", "attendee_handout"],
        capture_output=True, text=True, check=True,
    ).stdout
    handout_map = json.loads(handout_json)

    print(f"==> Writing {args.count} handouts to {out_dir} ...")
    for i in range(1, args.count + 1):
        num = f"{i:02d}"
        ns = f"attendee-{num}"
        sub_id = handout_map.get(ns, {}).get("apim_subscription_id", "")
        key = az_apim_key(ctx_static["RG"], ctx_static["APIM"], ns, sub_id)
        ctx = dict(ctx_static, APIM_KEY=key)
        path = out_dir / f"handout-{ns}.docx"
        write_handout_docx(path, ns, ctx)
        # Reset perms in case umask was too loose.
        try:
            path.chmod(0o600)
        except OSError:
            pass
        status = "ok" if key else "EMPTY KEY"
        print(f"    {ns:>12}  {path.name:>28}  ({status})")

    print("")
    print(f"Done. {args.count} files in {out_dir}")
    print("Fill in Name / Email / TAP per file inside Word before sending.")


if __name__ == "__main__":
    main()
